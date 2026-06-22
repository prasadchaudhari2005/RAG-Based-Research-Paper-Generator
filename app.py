import io
import os
import json
import logging
import uuid
from flask import Flask, request, jsonify, render_template, send_file, send_from_directory
from flask_cors import CORS
from groq import Groq
import google.generativeai as genai
from PIL import Image as PILImage
from paper_intelligence import extract_paper_metadata

# ------------------ ENV SETUP ------------------
from dotenv import load_dotenv
load_dotenv()

GROQ_API_KEY = os.getenv("GROQ_API_KEY")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

# Create metadata directory if not exists
os.makedirs("metadata_store", exist_ok=True)

# ------------------ RAG IMPORTS ------------------
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document

# ------------------ PDF PARSING ------------------
try:
    from pypdf import PdfReader
except ImportError:
    from PyPDF2 import PdfReader

# ------------------ DOCX EXPORT ------------------
from docx import Document as DocxDocument
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ------------------ LOGGING ------------------
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = Flask(__name__, template_folder="templates", static_folder="static")
CORS(app)

# ------------------ GLOBAL STATE ------------------
# For a local application, we can use simple global variables
VECTORSTORE = None
EXTRACTED_IMAGES = [] # List of dict: {"id": int, "label": str, "bytes": bytes}
GENERATED_PAPER = None
GENERATED_PAPER_TOPIC = None
UPLOADED_FILENAMES = []

# =========================================================
# PDF HELPERS
# =========================================================

def extract_text_from_pdf(file_bytes):
    file_like = io.BytesIO(file_bytes)
    reader = PdfReader(file_like)
    pages = []
    for p in reader.pages:
        t = p.extract_text()
        if t:
            pages.append(t)
    return "\n\n".join(pages)

def extract_images_from_pdf(file_bytes, min_width=300):
    images = []
    try:
        file_like = io.BytesIO(file_bytes)
        reader = PdfReader(file_like)
        for i, page in enumerate(reader.pages):
            if not hasattr(page, "images"):
                continue
            for img in page.images:
                try:
                    pil = PILImage.open(io.BytesIO(img.data))
                    if pil.mode not in ("RGB", "RGBA"):
                        pil = pil.convert("RGB")
                    if pil.size[0] < min_width:
                        continue
                    buf = io.BytesIO()
                    pil.save(buf, format="PNG")
                    buf.seek(0)
                    images.append((f"Page {i+1}", buf.getvalue()))
                except:
                    continue
    except Exception as e:
        logger.error(f"Image extraction error: {e}")
    return images

# =========================================================
# RAG CORE
# =========================================================

def build_vectorstore(files_data):
    docs = []
    splitter = RecursiveCharacterTextSplitter(chunk_size=800, chunk_overlap=150)

    for filename, file_bytes in files_data:
        text = extract_text_from_pdf(file_bytes)
        chunks = splitter.split_text(text)
        for c in chunks:
            docs.append(
                Document(
                    page_content=c,
                    metadata={"source": filename}
                )
            )

    embeddings = HuggingFaceEmbeddings(
        model_name="sentence-transformers/all-MiniLM-L6-v2"
    )
    return FAISS.from_documents(docs, embeddings)

def rag_retrieve(vectorstore, query, k=15):
    if not vectorstore:
        return ""
    docs = vectorstore.similarity_search(query, k=k)
    return "\n\n".join(d.page_content for d in docs)

# =========================================================
# GROQ ANALYSIS (GAPS + IDEAS)
# =========================================================

def groq_generate(client, prompt, model):
    res = client.chat.completions.create(
        messages=[{"role": "user", "content": prompt}],
        model=model,
        temperature=0.4,
        max_tokens=2000
    )
    return res.choices[0].message.content

# =========================================================
# GEMINI IEEE PAPER GENERATION (RAG-GROUNDED)
# =========================================================

def generate_ieee_paper(topic, literature_context, references_context):
    if not GEMINI_API_KEY:
        raise ValueError("GEMINI_API_KEY not found in environment")

    genai.configure(api_key=GEMINI_API_KEY)
    model = genai.GenerativeModel("gemini-2.5-flash")

    def run_prompt(prompt):
        response = model.generate_content(prompt)
        if not response or not response.text:
            raise RuntimeError("Gemini returned empty response")
        return response.text

    part1 = run_prompt(f"""
You are writing an IEEE research paper.

STRICT RULES:
- Use ONLY the provided context
- DO NOT invent references
- Follow IEEE format

TOPIC:
{topic}

===== LITERATURE CONTEXT START =====
{literature_context[:8000]}
===== LITERATURE CONTEXT END =====

Write:
Title
Abstract (200 words)
I. Introduction
""")

    part2 = run_prompt(f"""
Continue IEEE paper on topic "{topic}"

===== CONTEXT =====
{literature_context[:8000]}

Write:
II. Literature Survey
III. Proposed Methodology
""")

    part3 = run_prompt(f"""
Finish IEEE paper.

===== REFERENCES CONTEXT =====
{references_context[:6000]}

Write:
IV. Expected Results
V. Conclusion
VI. References (IEEE format, ONLY from context)
""")

    return "\n\n".join([part1, part2, part3])

# =========================================================
# DOCX EXPORT
# =========================================================

def create_docx_report(title, content, images=None):
    doc = DocxDocument()
    doc.add_heading(title, 0)

    for line in content.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith(("I.", "II.", "III.", "IV.", "V.", "VI.", "Abstract", "Title")):
            doc.add_heading(line, level=2)
        else:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if images:
        doc.add_page_break()
        doc.add_heading("Appendix", level=1)
        for lbl, img_bytes in images:
            img_io = io.BytesIO(img_bytes)
            doc.add_paragraph(lbl)
            doc.add_picture(img_io, width=Inches(5))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# =========================================================
# FLASK ROUTES
# =========================================================

@app.route("/")
def index():
    return render_template("index.html")

@app.route("/api/status", methods=["GET"])
def get_status():
    return jsonify({
        "indexed": VECTORSTORE is not None,
        "files": UPLOADED_FILENAMES
    })

@app.route("/api/upload", methods=["POST"])
def upload_files():
    global VECTORSTORE, EXTRACTED_IMAGES, UPLOADED_FILENAMES
    if "files" not in request.files:
        return jsonify({"error": "No files uploaded"}), 400
    
    files = request.files.getlist("files")
    if not files or files[0].filename == "":
        return jsonify({"error": "No files selected"}), 400

    files_data = []
    new_filenames = []
    new_images = []

    for f in files:
        if f.filename.endswith(".pdf"):
            file_bytes = f.read()
            files_data.append((f.filename, file_bytes))
            new_filenames.append(f.filename)
            
            # 1. Extract text and generate metadata using paper_intelligence
            try:
                pdf_text = extract_text_from_pdf(file_bytes)
                metadata = extract_paper_metadata(pdf_text)
                metadata_path = os.path.join("metadata_store", f"{f.filename}.json")
                with open(metadata_path, "w", encoding="utf-8") as meta_file:
                    json.dump(metadata, meta_file, indent=2)
                logger.info(f"Successfully saved intelligence metadata for {f.filename}")
            except Exception as ex:
                logger.error(f"Failed to generate intelligence metadata for {f.filename}: {ex}")

            # 2. Extract images
            extracted = extract_images_from_pdf(file_bytes)
            for label, img_bytes in extracted:
                new_images.append({
                    "id": len(EXTRACTED_IMAGES) + len(new_images),
                    "label": f"{f.filename} - {label}",
                    "bytes": img_bytes
                })

    if not files_data:
        return jsonify({"error": "No valid PDF files found"}), 400

    try:
        logger.info(f"Building RAG index for: {new_filenames}")
        new_store = build_vectorstore(files_data)
        if VECTORSTORE is None:
            VECTORSTORE = new_store
        else:
            # Merge vectorstores
            VECTORSTORE.merge_from(new_store)
        
        UPLOADED_FILENAMES.extend(new_filenames)
        EXTRACTED_IMAGES.extend(new_images)
        
        return jsonify({
            "success": True,
            "message": "Papers indexed successfully",
            "files": UPLOADED_FILENAMES,
            "image_count": len(EXTRACTED_IMAGES)
        })
    except Exception as e:
        logger.error(f"Error building index: {e}")
        return jsonify({"error": str(e)}), 500

@app.route("/api/metadata", methods=["GET"])
def get_all_metadata():
    os.makedirs("metadata_store", exist_ok=True)
    meta_list = {}
    for filename in os.listdir("metadata_store"):
        if filename.endswith(".json"):
            paper_name = filename[:-5]  # remove .json
            try:
                with open(os.path.join("metadata_store", filename), "r", encoding="utf-8") as f:
                    meta_list[paper_name] = json.load(f)
            except Exception as e:
                logger.error(f"Error reading metadata file {filename}: {e}")
    return jsonify(meta_list)

@app.route("/api/metadata/<paper_name>", methods=["GET"])
def get_paper_metadata(paper_name):
    # handle both with or without .pdf
    if paper_name.endswith(".pdf"):
        filename = f"{paper_name}.json"
    else:
        filename = f"{paper_name}.pdf.json"
        if not os.path.exists(os.path.join("metadata_store", filename)):
            filename = f"{paper_name}.json"
            
    filepath = os.path.join("metadata_store", filename)
    if not os.path.exists(filepath):
        return jsonify({"error": "Metadata not found"}), 404
        
    try:
        with open(filepath, "r", encoding="utf-8") as f:
            return jsonify(json.load(f))
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/review", methods=["POST"])
def generate_review():
    global VECTORSTORE
    if not VECTORSTORE:
        return jsonify({"error": "No papers uploaded yet"}), 400
    
    if not GROQ_API_KEY:
        return jsonify({"error": "GROQ_API_KEY missing in environment"}), 500
    
    data = request.json or {}
    model_choice = data.get("model", "llama-3.3-70b-versatile")
    
    try:
        client = Groq(api_key=GROQ_API_KEY)
        ctx = rag_retrieve(
            VECTORSTORE,
            "Summarize problem statements, methods, datasets, and metrics used across all papers."
        )
        review = groq_generate(
            client,
            f"Generate a unified literature review:\n{ctx}",
            model_choice
        )
        return jsonify({"review": review})
    except Exception as e:
        logger.error(f"Error generating review: {e}")
        return jsonify({"error": str(e)}), 500

@app.route("/api/gaps", methods=["POST"])
def generate_gaps():
    global VECTORSTORE
    if not VECTORSTORE:
        return jsonify({"error": "No papers uploaded yet"}), 400
    
    if not GROQ_API_KEY:
        return jsonify({"error": "GROQ_API_KEY missing in environment"}), 500
    
    data = request.json or {}
    model_choice = data.get("model", "llama-3.3-70b-versatile")
    
    try:
        client = Groq(api_key=GROQ_API_KEY)
        ctx = rag_retrieve(
            VECTORSTORE,
            "Identify limitations, future work, and unexplored research directions."
        )
        gaps = groq_generate(
            client,
            f"From the following context:\n{ctx}\n\nProvide:\n1. Research gaps\n2. 3 Novel research ideas with titles and short descriptions",
            model_choice
        )
        return jsonify({"gaps": gaps})
    except Exception as e:
        logger.error(f"Error generating gaps: {e}")
        return jsonify({"error": str(e)}), 500

@app.route("/api/generate-paper", methods=["POST"])
def api_generate_paper():
    global VECTORSTORE, GENERATED_PAPER, GENERATED_PAPER_TOPIC
    if not VECTORSTORE:
        return jsonify({"error": "No papers uploaded yet"}), 400
    
    if not GEMINI_API_KEY:
        return jsonify({"error": "GEMINI_API_KEY missing in environment"}), 500
    
    data = request.json or {}
    topic = data.get("topic")
    if not topic:
        return jsonify({"error": "Topic is required"}), 400
    
    try:
        literature_ctx = rag_retrieve(
            VECTORSTORE,
            f"Provide key literature insights relevant to the selected topic: {topic}"
        )
        references_ctx = rag_retrieve(
            VECTORSTORE,
            "Extract references including authors, title, venue, and year."
        )
        
        paper = generate_ieee_paper(topic, literature_ctx, references_ctx)
        GENERATED_PAPER = paper
        GENERATED_PAPER_TOPIC = topic
        
        return jsonify({"paper": paper})
    except Exception as e:
        logger.error(f"Error generating paper: {e}")
        return jsonify({"error": str(e)}), 500

@app.route("/api/ask", methods=["POST"])
def api_ask():
    global VECTORSTORE
    if not VECTORSTORE:
        return jsonify({"error": "No papers uploaded yet"}), 400
    
    if not GROQ_API_KEY:
        return jsonify({"error": "GROQ_API_KEY missing in environment"}), 500
    
    data = request.json or {}
    question = data.get("question")
    model_choice = data.get("model", "llama-3.3-70b-versatile")
    if not question:
        return jsonify({"error": "Question is required"}), 400
        
    try:
        client = Groq(api_key=GROQ_API_KEY)
        ctx = rag_retrieve(VECTORSTORE, question)
        ans = groq_generate(
            client,
            f"Based on the following context, answer the user's question:\nContext: {ctx}\n\nQuestion: {question}",
            model_choice
        )
        return jsonify({"answer": ans})
    except Exception as e:
        logger.error(f"Error answering question: {e}")
        return jsonify({"error": str(e)}), 500

@app.route("/api/images", methods=["GET"])
def get_images():
    return jsonify([
        {"id": img["id"], "label": img["label"]}
        for img in EXTRACTED_IMAGES
    ])

@app.route("/api/image/<int:img_id>", methods=["GET"])
def serve_image(img_id):
    for img in EXTRACTED_IMAGES:
        if img["id"] == img_id:
            return send_file(
                io.BytesIO(img["bytes"]),
                mimetype="image/png"
            )
    return jsonify({"error": "Image not found"}), 404

@app.route("/api/download-docx", methods=["GET"])
def download_docx():
    global GENERATED_PAPER, GENERATED_PAPER_TOPIC
    if not GENERATED_PAPER:
        return jsonify({"error": "No paper has been generated yet"}), 400
    
    try:
        # Prepare list of extracted images if we want to include them in the appendix
        # For simplicity, we can pass all extracted images to the docx generator or none.
        # Let's pass the images in memory
        images_to_append = [(img["label"], img["bytes"]) for img in EXTRACTED_IMAGES]
        
        docx_buf = create_docx_report(
            GENERATED_PAPER_TOPIC or "Generated IEEE Research Paper",
            GENERATED_PAPER,
            images=images_to_append
        )
        
        filename = f"IEEE_Paper_{uuid.uuid4().hex[:8]}.docx"
        return send_file(
            docx_buf,
            as_attachment=True,
            download_name=filename,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as e:
        logger.error(f"Error downloading docx: {e}")
        return jsonify({"error": str(e)}), 500

# =========================================================
# MAIN
# =========================================================

if __name__ == "__main__":
    # Check for keys
    if not GROQ_API_KEY:
        logger.warning("GROQ_API_KEY missing in env")
    if not GEMINI_API_KEY:
        logger.warning("GEMINI_API_KEY missing in env")
        
    app.run(host="0.0.0.0", port=5000, debug=True)
